"""Extração de texto puro de documentos enviados (.txt, .md, .docx, .pdf).

Usado pela importação de copy: o usuário pode subir um documento em vez de
colar o texto. O texto extraído alimenta a extração estruturada por IA.
"""

from __future__ import annotations

import io


def extract_text_from_upload(filename: str, data: bytes) -> str:
    """Devolve o texto puro de um arquivo. Lança RuntimeError em formato não suportado."""
    name = (filename or "").lower()

    if name.endswith((".txt", ".md", ".markdown")):
        return _decode(data)

    if name.endswith(".docx"):
        return _from_docx(data)

    if name.endswith(".pdf"):
        return _from_pdf(data)

    # fallback: tenta decodificar como texto
    try:
        return _decode(data)
    except Exception as exc:  # noqa: BLE001
        raise RuntimeError(f"Formato não suportado: {filename}") from exc


def _decode(data: bytes) -> str:
    for enc in ("utf-8", "latin-1"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="ignore")


def _from_docx(data: bytes) -> str:
    try:
        from docx import Document  # python-docx
    except ImportError as exc:
        raise RuntimeError("python-docx não instalado no servidor") from exc

    doc = Document(io.BytesIO(data))
    lines: list[str] = [p.text for p in doc.paragraphs]
    # também extrai texto de tabelas (briefings costumam usar tabelas)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                lines.append(" | ".join(cells))
    text = "\n".join(lines).strip()
    if not text:
        raise RuntimeError("Documento .docx sem texto extraível")
    return text


def _from_pdf(data: bytes) -> str:
    try:
        from pdfminer.high_level import extract_text  # pdfminer.six
    except ImportError as exc:
        raise RuntimeError("pdfminer.six não instalado no servidor") from exc

    text = (extract_text(io.BytesIO(data)) or "").strip()
    if not text:
        raise RuntimeError("PDF sem texto extraível (pode ser digitalizado/imagem)")
    return text
